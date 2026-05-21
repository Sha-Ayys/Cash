from flask import Flask, render_template, request, jsonify, send_file
import json, os, io
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
DATA_FILE = "data.json"

def load_data():
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE, "r") as f:
            return json.load(f)
    return {"products": [], "transactions": [], "qris_config": {"image": "", "merchant_name": "Toko Saya"}}

def save_data(data):
    with open(DATA_FILE, "w") as f:
        json.dump(data, f, indent=2)

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/api/products", methods=["GET"])
def get_products():
    data = load_data()
    return jsonify(data["products"])

@app.route("/api/products", methods=["POST"])
def add_product():
    data = load_data()
    body = request.json
    product = {
        "id": int(datetime.now().timestamp() * 1000),
        "name": body["name"],
        "price": float(body["price"]),
        "stock": int(body.get("stock", 0)),
        "category": body.get("category", "Umum"),
        "sold": 0
    }
    data["products"].append(product)
    save_data(data)
    return jsonify(product), 201

@app.route("/api/products/<int:pid>", methods=["PUT"])
def update_product(pid):
    data = load_data()
    for p in data["products"]:
        if p["id"] == pid:
            p["name"] = request.json.get("name", p["name"])
            p["price"] = float(request.json.get("price", p["price"]))
            p["stock"] = int(request.json.get("stock", p["stock"]))
            p["category"] = request.json.get("category", p.get("category", "Umum"))
            save_data(data)
            return jsonify(p)
    return jsonify({"error": "Not found"}), 404

@app.route("/api/products/<int:pid>", methods=["DELETE"])
def delete_product(pid):
    data = load_data()
    data["products"] = [p for p in data["products"] if p["id"] != pid]
    save_data(data)
    return jsonify({"ok": True})

@app.route("/api/transactions", methods=["GET"])
def get_transactions():
    data = load_data()
    return jsonify(data["transactions"])

@app.route("/api/transactions", methods=["POST"])
def add_transaction():
    data = load_data()
    body = request.json
    trx = {
        "id": int(datetime.now().timestamp() * 1000),
        "items": body["items"],
        "total": body["total"],
        "payment_method": body["payment_method"],
        "amount_paid": body.get("amount_paid", body["total"]),
        "change": body.get("change", 0),
        "date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "cashier": body.get("cashier", "Kasir")
    }
    for item in body["items"]:
        for p in data["products"]:
            if p["id"] == item["id"]:
                p["sold"] = p.get("sold", 0) + item["qty"]
                p["stock"] = max(0, p.get("stock", 0) - item["qty"])
    data["transactions"].append(trx)
    save_data(data)
    return jsonify(trx), 201

@app.route("/api/qris", methods=["GET", "POST"])
def qris_config():
    data = load_data()
    if request.method == "POST":
        body = request.json
        data["qris_config"] = {
            "image": body.get("image", data.get("qris_config", {}).get("image", "")),
            "merchant_name": body.get("merchant_name", "Toko Saya")
        }
        save_data(data)
        return jsonify(data["qris_config"])
    return jsonify(data.get("qris_config", {"image": "", "merchant_name": "Toko Saya"}))

@app.route("/api/export/excel")
def export_excel():
    data = load_data()
    wb = Workbook()
    ws_prod = wb.active
    ws_prod.title = "Produk"
    ws_trx = wb.create_sheet("Transaksi")
    ws_summary = wb.create_sheet("Ringkasan")

    header_fill = PatternFill("solid", fgColor="1a1a2e")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    alt_fill = PatternFill("solid", fgColor="F0F4FF")
    center = Alignment(horizontal="center", vertical="center")
    thin_border = Border(
        left=Side(style='thin', color='CCCCCC'),
        right=Side(style='thin', color='CCCCCC'),
        top=Side(style='thin', color='CCCCCC'),
        bottom=Side(style='thin', color='CCCCCC')
    )

    ws_prod.append(["ID", "Nama Produk", "Kategori", "Harga (Rp)", "Stok", "Terjual", "Pendapatan (Rp)"])
    for i, cell in enumerate(ws_prod[1]):
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center
        cell.border = thin_border

    ws_prod.column_dimensions['A'].width = 15
    ws_prod.column_dimensions['B'].width = 25
    ws_prod.column_dimensions['C'].width = 15
    ws_prod.column_dimensions['D'].width = 18
    ws_prod.column_dimensions['E'].width = 10
    ws_prod.column_dimensions['F'].width = 10
    ws_prod.column_dimensions['G'].width = 20

    for idx, p in enumerate(data["products"]):
        revenue = p.get("sold", 0) * p["price"]
        row = [p["id"], p["name"], p.get("category", "Umum"),
               p["price"], p.get("stock", 0), p.get("sold", 0), revenue]
        ws_prod.append(row)
        if idx % 2 == 1:
            for cell in ws_prod[idx + 2]:
                cell.fill = alt_fill
        for cell in ws_prod[idx + 2]:
            cell.border = thin_border
            cell.alignment = center

    ws_trx.append(["ID", "Tanggal", "Kasir", "Metode Bayar", "Total (Rp)", "Dibayar (Rp)", "Kembalian (Rp)", "Item"])
    for i, cell in enumerate(ws_trx[1]):
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center
        cell.border = thin_border

    for col in ['A','B','C','D','E','F','G','H']:
        ws_trx.column_dimensions[col].width = 20

    for idx, t in enumerate(data["transactions"]):
        items_str = ", ".join([f"{i['name']} x{i['qty']}" for i in t["items"]])
        row = [t["id"], t["date"], t.get("cashier",""), t["payment_method"],
               t["total"], t.get("amount_paid", t["total"]), t.get("change", 0), items_str]
        ws_trx.append(row)
        if idx % 2 == 1:
            for cell in ws_trx[idx + 2]:
                cell.fill = alt_fill
        for cell in ws_trx[idx + 2]:
            cell.border = thin_border

    total_rev = sum(t["total"] for t in data["transactions"])
    total_items = sum(sum(i["qty"] for i in t["items"]) for t in data["transactions"])
    ws_summary.append(["RINGKASAN PENJUALAN"])
    ws_summary["A1"].font = Font(bold=True, size=16, color="1a1a2e")
    ws_summary.append(["Tanggal Ekspor", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
    ws_summary.append(["Total Transaksi", len(data["transactions"])])
    ws_summary.append(["Total Item Terjual", total_items])
    ws_summary.append(["Total Pendapatan", f"Rp {total_rev:,.0f}"])
    ws_summary.column_dimensions['A'].width = 25
    ws_summary.column_dimensions['B'].width = 30

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = f"laporan_kasir_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                     as_attachment=True, download_name=fname)

@app.route("/api/export/word")
def export_word():
    data = load_data()
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'

    title = doc.add_heading('Laporan Kasir', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    doc.add_paragraph(f"Tanggal: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    doc.add_paragraph(f"Total Transaksi: {len(data['transactions'])}")
    total_rev = sum(t['total'] for t in data['transactions'])
    doc.add_paragraph(f"Total Pendapatan: Rp {total_rev:,.0f}")
    doc.add_paragraph("")

    doc.add_heading('Daftar Produk', level=1)
    if data["products"]:
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        headers = ['Nama', 'Kategori', 'Harga', 'Stok', 'Terjual', 'Pendapatan']
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].font.bold = True
        for p in data['products']:
            row = table.add_row().cells
            row[0].text = p['name']
            row[1].text = p.get('category', 'Umum')
            row[2].text = f"Rp {p['price']:,.0f}"
            row[3].text = str(p.get('stock', 0))
            row[4].text = str(p.get('sold', 0))
            row[5].text = f"Rp {p.get('sold',0)*p['price']:,.0f}"

    doc.add_paragraph("")
    doc.add_heading('Riwayat Transaksi', level=1)
    if data["transactions"]:
        table2 = doc.add_table(rows=1, cols=5)
        table2.style = 'Table Grid'
        hdr2 = table2.rows[0].cells
        headers2 = ['Tanggal', 'Metode', 'Total', 'Dibayar', 'Kembalian']
        for i, h in enumerate(headers2):
            hdr2[i].text = h
            hdr2[i].paragraphs[0].runs[0].font.bold = True
        for t in data['transactions']:
            row = table2.add_row().cells
            row[0].text = t['date']
            row[1].text = t['payment_method']
            row[2].text = f"Rp {t['total']:,.0f}"
            row[3].text = f"Rp {t.get('amount_paid', t['total']):,.0f}"
            row[4].text = f"Rp {t.get('change', 0):,.0f}"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f"laporan_kasir_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name=fname)

if __name__ == "__main__":
    app.run(debug=True, port=5000)
